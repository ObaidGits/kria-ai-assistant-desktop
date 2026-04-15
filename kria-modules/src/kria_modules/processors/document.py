"""
Document processor — Pre-Cognitive document extraction.

Extracts text, structure, and metadata from PDF, DOCX, CSV, and
plain-text files before the LLM processes them. Tier-aware depth.
"""

import logging
import os
from pathlib import Path
from typing import Any

logger = logging.getLogger("kria.processors.document")

METHODS = ["extract"]

# Tier → max chars extracted
_MAX_CHARS = {"lite": 4_000, "standard": 16_000, "performance": 64_000, "high": 256_000}


def extract(params: dict) -> dict:
    """
    Extract content from a document file.

    Params:
        file_path: str — path to document
        max_chars: int — override char budget (optional)
    """
    file_path = params.get("file_path", "")
    if not file_path or not os.path.isfile(file_path):
        raise FileNotFoundError(file_path)

    tier = params.get("_tier", "standard")
    max_chars = params.get("max_chars", _MAX_CHARS.get(tier, 16_000))
    ext = Path(file_path).suffix.lower()

    result: dict[str, Any] = {
        "file_path": file_path,
        "file_type": ext.lstrip("."),
        "size_kb": round(os.path.getsize(file_path) / 1024, 1),
    }

    if ext == ".pdf":
        result.update(_extract_pdf(file_path, max_chars, tier))
    elif ext in (".docx", ".doc"):
        result.update(_extract_docx(file_path, max_chars))
    elif ext in (".csv", ".tsv"):
        result.update(_extract_csv(file_path, max_chars, ext))
    elif ext in (".txt", ".md", ".log", ".json", ".yaml", ".yml", ".toml", ".xml", ".html"):
        result.update(_extract_text(file_path, max_chars))
    else:
        result["error"] = f"Unsupported extension: {ext}"

    # Truncation marker
    text = result.get("text", "")
    if len(text) > max_chars:
        result["text"] = text[:max_chars] + "\n[TRUNCATED]"
        result["truncated"] = True
    else:
        result["truncated"] = False

    result["char_count"] = len(result.get("text", ""))

    # Summary
    lines = result.get("text", "").split("\n")
    preview = "\n".join(lines[:5])[:300]
    result["summary"] = f"{result['file_type'].upper()} | {result['size_kb']}KB | {result['char_count']} chars | Preview: {preview}"

    return result


def _extract_pdf(path: str, max_chars: int, tier: str) -> dict:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    meta: dict[str, Any] = {
        "page_count": len(doc),
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
    }

    pages = []
    total = 0
    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append({"page": i + 1, "text": text})
        total += len(text)
        if total >= max_chars:
            break

    full_text = "\n\n".join(p["text"] for p in pages)

    # If high tier, also extract tables
    tables = []
    if tier in ("performance", "high"):
        for i, page in enumerate(doc):
            try:
                ts = page.find_tables()
                for t in ts:
                    rows = t.extract()
                    if rows:
                        tables.append({"page": i + 1, "rows": rows[:20]})
            except Exception:
                pass
            if i >= 10:
                break

    result = {"text": full_text, "metadata": meta, "pages_extracted": len(pages)}
    if tables:
        result["tables"] = tables
    return result


def _extract_docx(path: str, max_chars: int) -> dict:
    from docx import Document

    doc = Document(path)
    paragraphs = []
    total = 0
    for para in doc.paragraphs:
        text = para.text
        paragraphs.append(text)
        total += len(text)
        if total >= max_chars:
            break

    # Extract tables
    tables = []
    for table in doc.tables[:5]:
        rows = []
        for row in table.rows[:20]:
            rows.append([cell.text for cell in row.cells])
        if rows:
            tables.append({"rows": rows})

    result: dict[str, Any] = {
        "text": "\n".join(paragraphs),
        "paragraph_count": len(doc.paragraphs),
    }
    if tables:
        result["tables"] = tables
    return result


def _extract_csv(path: str, max_chars: int, ext: str) -> dict:
    import pandas as pd

    sep = "\t" if ext == ".tsv" else ","
    try:
        df = pd.read_csv(path, sep=sep, nrows=500)
    except Exception as e:
        return {"text": "", "error": f"CSV parse error: {e}"}

    info: dict[str, Any] = {
        "columns": list(df.columns),
        "shape": list(df.shape),
        "dtypes": {col: str(dt) for col, dt in df.dtypes.items()},
    }

    text = df.head(50).to_string(index=False)
    if len(text) > max_chars:
        text = text[:max_chars]

    return {"text": text, "table_info": info}


def _extract_text(path: str, max_chars: int) -> dict:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read(max_chars + 1)

    return {"text": text}
